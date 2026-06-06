#!/usr/bin/env python3
"""Prepend a CUFE cover page image to a DOCX without rebuilding the cover in Word."""

from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
VENDOR = ROOT / "vendor"
SCRIPT_DIR = Path(__file__).resolve().parent


def enable_vendor_if_compatible() -> bool:
    if os.environ.get("CUFE_COVER_NO_VENDOR"):
        return False
    if not VENDOR.exists() or str(VENDOR) in sys.path:
        return VENDOR.exists()
    if not os.environ.get("CUFE_COVER_FORCE_VENDOR"):
        cache_tag = sys.implementation.cache_tag
        binary_modules = list(VENDOR.rglob("*.so")) + list(VENDOR.rglob("*.pyd"))
        incompatible = [
            path.name
            for path in binary_modules
            if "cpython-" in path.name and cache_tag not in path.name
        ]
        if incompatible:
            return False
    sys.path.insert(0, str(VENDOR))
    return True


def dependency_error(package: str, exc: BaseException) -> SystemExit:
    return SystemExit(
        f"Missing dependency: {package}. Install with `python3 -m pip install -r requirements.txt`, "
        f"use a matching Python for cufe-report-cover/vendor, or set CUFE_COVER_NO_VENDOR=1 "
        f"to ignore vendored packages. Current Python ABI is {sys.implementation.cache_tag}. "
        f"Original error: {exc}"
    )


def import_docx():
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.shared import Cm, Pt
        return Document, WD_SECTION, Cm, Pt
    except Exception as first_exc:
        if enable_vendor_if_compatible():
            try:
                from docx import Document
                from docx.enum.section import WD_SECTION
                from docx.shared import Cm, Pt
                return Document, WD_SECTION, Cm, Pt
            except Exception as vendor_exc:
                raise dependency_error("python-docx", vendor_exc) from vendor_exc
        raise dependency_error("python-docx", first_exc) from first_exc


Document, WD_SECTION, Cm, Pt = import_docx()


def insert_paragraph_at_start(document):
    body = document._body._element
    paragraph = document.add_paragraph()
    body.remove(paragraph._element)
    body.insert(0, paragraph._element)
    return paragraph


def render_cover_png(args, output_png: Path) -> None:
    cmd = [
        sys.executable,
        str(SCRIPT_DIR / "render_cover_png.py"),
        "--type",
        args.type,
        "--output",
        str(output_png),
        "--dpi",
        str(args.dpi),
        "--font-index",
        str(args.font_index),
        "--field-offset-x",
        str(args.field_offset_x),
        "--field-offset-y",
        str(args.field_offset_y),
    ]
    if args.data:
        cmd.extend(["--data", args.data])
    if args.font_path:
        cmd.extend(["--font-path", args.font_path])
    for field in args.field:
        cmd.extend(["--field", field])
    subprocess.run(cmd, check=True)


def prepend_cover(input_docx: Path, output_docx: Path, cover_png: Path, body_margins_cm: tuple[float, float, float, float]) -> None:
    document = Document(str(input_docx))

    cover_section = document.sections[0]
    cover_section.page_width = Cm(21.0)
    cover_section.page_height = Cm(29.7)
    cover_section.top_margin = Cm(0)
    cover_section.bottom_margin = Cm(0)
    cover_section.left_margin = Cm(0)
    cover_section.right_margin = Cm(0)
    cover_section.header_distance = Cm(0)
    cover_section.footer_distance = Cm(0)

    cover_paragraph = insert_paragraph_at_start(document)
    cover_paragraph.alignment = 1
    cover_paragraph.paragraph_format.space_before = Pt(0)
    cover_paragraph.paragraph_format.space_after = Pt(0)
    cover_paragraph.paragraph_format.line_spacing = 1
    cover_run = cover_paragraph.add_run()
    cover_run.add_picture(str(cover_png), width=Cm(21.0))

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    top, bottom, left, right = body_margins_cm
    body_section.top_margin = Cm(top)
    body_section.bottom_margin = Cm(bottom)
    body_section.left_margin = Cm(left)
    body_section.right_margin = Cm(right)

    body = document._body._element
    section_break = body[-2]
    body.remove(section_break)
    body.insert(1, section_break)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))


def main() -> int:
    parser = argparse.ArgumentParser(description="Prepend a rendered CUFE cover image to a DOCX.")
    parser.add_argument("--input-docx", required=True, help="Existing DOCX body.")
    parser.add_argument("--output-docx", required=True, help="Output DOCX with the cover prepended.")
    parser.add_argument("--type", required=True, choices=["course-paper", "lab-report"], help="Cover template to use.")
    parser.add_argument("--data", help="YAML or JSON file containing cover fields.")
    parser.add_argument("--field", action="append", default=[], help="Override/add one field as key=value. Can be repeated.")
    parser.add_argument("--cover-png", help="Use an existing rendered cover PNG instead of generating one.")
    parser.add_argument("--preview-png", help="Where to write or copy the generated cover PNG for inspection.")
    parser.add_argument("--work-dir", help="Directory for temporary cover render artifacts.")
    parser.add_argument("--dpi", type=int, default=220, help="Render DPI when generating the cover PNG.")
    parser.add_argument("--font-path", help="CJK font path for field values.")
    parser.add_argument("--font-index", type=int, default=2, help="Font face index for TTC collections; Noto Sans CJK SC is usually 2.")
    parser.add_argument("--field-offset-x", type=float, default=0, help="Global field x-offset in output pixels.")
    parser.add_argument("--field-offset-y", type=float, default=0, help="Global field y-offset in output pixels.")
    parser.add_argument("--body-margin-top", type=float, default=2.5, help="Body top margin in cm.")
    parser.add_argument("--body-margin-bottom", type=float, default=2.5, help="Body bottom margin in cm.")
    parser.add_argument("--body-margin-left", type=float, default=2.3, help="Body left margin in cm.")
    parser.add_argument("--body-margin-right", type=float, default=2.3, help="Body right margin in cm.")
    args = parser.parse_args()

    input_docx = Path(args.input_docx)
    output_docx = Path(args.output_docx)
    body_margins = (args.body_margin_top, args.body_margin_bottom, args.body_margin_left, args.body_margin_right)

    if args.cover_png:
        cover_png = Path(args.cover_png)
    else:
        if args.work_dir:
            work_dir = Path(args.work_dir)
            work_dir.mkdir(parents=True, exist_ok=True)
            cover_png = work_dir / "cover.png"
            render_cover_png(args, cover_png)
        else:
            with tempfile.TemporaryDirectory() as tmp:
                cover_png = Path(tmp) / "cover.png"
                render_cover_png(args, cover_png)
                if args.preview_png:
                    preview = Path(args.preview_png)
                    preview.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(cover_png, preview)
                prepend_cover(input_docx, output_docx, cover_png, body_margins)
                print(output_docx)
                return 0

    if args.preview_png:
        preview = Path(args.preview_png)
        preview.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(cover_png, preview)
    prepend_cover(input_docx, output_docx, cover_png, body_margins)
    print(output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
